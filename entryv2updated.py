from Tkinter import*
import Tkinter as Tkinter
import MySQLdb
import tkFont
from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# subject1="Computer Networks"
# subject2="Operating Systems"
# subject3="Scripting Languages"
# subject4="Software Engineering"
# subject5="Human Computer Interaction"



#Edit Interface
def editQuestions():
	editRoot = Tk()
	editRoot.title("Edit Interface")
	editRoot.configure(bg = "#D3D3D3")

	def cancel():
		editRoot.destroy()
	def confirmQuestion():
		db=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
		cursor=db.cursor()
		editChangesSQL="UPDATE qBank SET question="+e1f1.get()+" ,opt1="+e2f1.get()+" , opt2="+e3f1.get()+" , opt3="+e4f1.get()+" , opt4="+f5f1.get()+" , answer="+editCorrectAnswer.get()+" , subject="+editSubjectVal.get() +" WHERE qId="+e1f0.get()
		try:
			cursor.execute(editChangesSQL)
			db.commit()
		except Exception, e:
			#raise e
			db.rollback()
			print "Error Updating the Database. Please Try later."
		db.close()

	def search():
		db=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
		cursor=db.cursor()
		editQuestionId=e1f0.get()
		ShowEditQuestionSQL="SELECT * FROM qBank WHERE qId="+editQuestionId
		cursor.execute(ShowEditQuestionSQL)
		editQues=cursor.fetchone()
		print editQues
		e1f1.delete(0,END)
		e1f1.insert(0,editQues[1])
		e2f1.delete(0,END)
		e2f1.insert(0,editQues[3])
		e3f1.delete(0,END)
		e3f1.insert(0,editQues[4])
		e4f1.delete(0,END)
		e4f1.insert(0,editQues[5])
		e5f1.delete(0,END)
		e5f1.insert(0,editQues[6])
		if(editQues[7]==1):
			rb6f1.select()
			rb7f1.deselect()
			rb8f1.deselect()
			rb9f1.deselect()
		elif(editQues[7]==2):
			rb6f1.deselect()
			rb7f1.select()
			rb8f1.deselect()
			rb9f1.deselect()
		elif(editQues[7]==3):
			rb6f1.deselect()
			rb7f1.deselect()
			rb8f1.select()
			rb9f1.deselect()
		elif(editQues[7]==4):
			rb6f1.deselect()
			rb7f1.deselect()
			rb8f1.deselect()
			rb9f1.select()
		#print eom1f1.index(subjectList[2])
		editSubjectVal.set(editQues[2])
		eom1f1.update()


	ws = editRoot.winfo_screenwidth()
	hs = editRoot.winfo_screenheight()

	w=1000
	h=600
	x = (ws/2) - (w/2)
	y = (hs/2) - (h/2)
	editRoot.geometry('%dx%d+%d+%d' % (w, h, x, y))
	# editscrndim=str(ws/2)+'x'+str(hs/2)
	# editoot.geometry('800x400')
	editSubjectVal=StringVar()
	editCorrectAnswer=StringVar()
	labelframe0=LabelFrame(editRoot,text="Enter Question Number ",bg="#D3D3D3")
	labelframe0.place(relx=0.02,rely=0.08,relwidth=0.95,relheight=0.1)

	l1f0=Label(labelframe0 ,bg="#D3D3D3", text="Enter Question Number :" )
	l1f0.place(relx=0.013,rely=0.1)
	e1f0=Entry(labelframe0 ,highlightbackground="#1E90FF", width = 20)
	e1f0.place(relx=0.218, rely=0.1 )

	searchButton=Button(labelframe0,text="Search",activebackground="#1E90FF",highlightbackground="#1E90FF", width=20, bg="#87CEFA",command=search)
	searchButton.place(rely=0.02,relx=0.425)

	labelframe1=LabelFrame(editRoot,text="Edit Interface",bg="#D3D3D3")
	labelframe1.place(relx=0.02,rely=0.2,relwidth=0.95,relheight=0.7)

	
	
	l1f1=Label(labelframe1 ,bg="#D3D3D3", text="Selected Question :" , height = 2 , width = 15)
	l1f1.grid(row=0 , column = 0 ,pady=10 ,sticky=W, padx=12)

	e1f1=Entry(labelframe1 ,highlightbackground="#1E90FF", width = 90)
	e1f1.grid(row=0 , column=1 , sticky=W,columnspan=4)

	l2f1=Label(labelframe1,bg="#D3D3D3",text="Enter the Options :")
	l2f1.grid(row=1 , column=0,sticky=W, padx=12)

	l3f1=Label(labelframe1,bg="#D3D3D3",text="Option A ------------->")
	l3f1.grid(row=2 , column=0 , sticky=W , padx=12)

	e2f1=Entry(labelframe1,highlightbackground="#1E90FF",width = 50)
	e2f1.grid(row=2 , column=1 , sticky=W)

	l4f1=Label(labelframe1,bg="#D3D3D3",text="Option B ------------->")
	l4f1.grid(row=3 , column=0 , sticky=W , padx=12)

	e3f1=Entry(labelframe1,highlightbackground="#1E90FF",width = 50)
	e3f1.grid(row=3 , column=1 , sticky=W)

	l5f1=Label(labelframe1,bg="#D3D3D3",text="Option C ------------->")
	l5f1.grid(row=4 , column=0 , sticky=W , padx=12)

	e4f1=Entry(labelframe1,highlightbackground="#1E90FF",width = 50)
	e4f1.grid(row=4 , column=1 , sticky=W)

	l6f1=Label(labelframe1,bg="#D3D3D3",text="Option D ------------->")
	l6f1.grid(row=5 , column=0 , sticky=W , padx=12)

	e5f1=Entry(labelframe1,highlightbackground="#1E90FF",width = 50)
	e5f1.grid(row=5 , column=1 ,  sticky=W)

	l7f1=Label(labelframe1,bg="#D3D3D3",text="Choose Subject :")
	l7f1.grid(row=7 , column = 0 , padx=12 , sticky=W)

	# rb1f1=Radiobutton(labelframe1,activebackground="#87CEFA", selectcolor="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",width=35,text="Computer Networks" , variable=subjectvalinf1,value="Computer Networks" , indicatoron=0)
	# rb1f1.grid(row=7,column = 1,sticky=W)

	# rb2f1=Radiobutton(labelframe1,activebackground="#87CEFA",selectcolor="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",width=35,text="Operating Systems" , variable=subjectvalinf1,value="Operating Systems", indicatoron=0)
	# rb2f1.grid(row=8,column = 1,sticky=W)

	# rb3f1=Radiobutton(labelframe1,activebackground="#87CEFA",selectcolor="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",width=35,text="Scripting Languages" , variable=subjectvalinf1,value="Scripting Languages", indicatoron=0)
	# rb3f1.grid(row=9,column = 1,sticky=W)

	# rb4f1=Radiobutton(labelframe1,activebackground="#87CEFA",selectcolor="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",width=35,text="Software Engineering" , variable=subjectvalinf1,value="Software Engineering", indicatoron=0)
	# rb4f1.grid(row=10,column = 1,sticky=W)

	# rb5f1=Radiobutton(labelframe1,activebackground="#87CEFA",selectcolor="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",width=35,text="Human Computer Interaction" , variable=subjectvalinf1,value="Human Computer Interaction", indicatoron=0)
	# rb5f1.grid(row=11,column = 1,sticky=W)

	eom1f1 = OptionMenu(labelframe1, editSubjectVal, *subjectList)
	eom1f1.config(width=15 , activebackground="#1E90FF")
	print subjectList[3]
	editSubjectVal.set(subjectList[3])
	eom1f1.grid(row = 7 , column =1 ,sticky=W)

	l8f1=Label(labelframe1,bg="#D3D3D3",text="Choose the correct answer ->")
	l8f1.grid(row=6 , column=0 , padx=12)

	rb6f1=Radiobutton(labelframe1,activebackground="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",text="A",variable=editCorrectAnswer,value=1)
	rb6f1.grid(row=6,column=1,sticky=W)

	rb7f1=Radiobutton(labelframe1,activebackground="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",text="B",variable=editCorrectAnswer,value=2)
	rb7f1.grid(row=6,column=1,sticky=W,padx=40)

	rb8f1=Radiobutton(labelframe1,activebackground="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",text="C",variable=editCorrectAnswer,value=3)
	rb8f1.grid(row=6,column=1,sticky=W,padx=80)

	rb9f1=Radiobutton(labelframe1,activebackground="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",text="D",variable=editCorrectAnswer,value=4)
	rb9f1.grid(row=6,column=1,sticky=W,padx=120)

	confirmbutton=Button(labelframe1,text="CONFIRM CHANGES",activebackground="#1E90FF",highlightbackground="#1E90FF",command=confirmQuestion , width=20 , bg="#87CEFA")
	confirmbutton.grid(row=12 ,column=1 ,sticky=W  , pady = 10)

	cancelButton=Button(labelframe1,text="CANCEL",activebackground="#1E90FF",highlightbackground="#1E90FF",command=cancel , width=20, bg
		="#87CEFA")
	cancelButton.grid(row=12,column=1,sticky=E,padx=100)
	


	editRoot.mainloop()

#--------------------------Main Entry Interface------------------------
root = Tk()
root.title("Question Paper Generator v1.02")
root.configure(bg = "#D3D3D3")
generatedQPaper=0
subjectList=[]
optList=[]
subjectQuesCount={}
def insertQuestion():
	#
	db=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	cursor=db.cursor()
	sql="INSERT INTO qBank(`question`,`subject`,`opt1`,`opt2`,`opt3`,`opt4`,`answer`) VALUES('"+e1f1.get()+"','"+subjectvalinf1.get()+"','"+e2f1.get()+"','"+e3f1.get()+"','"+e4f1.get()+"','"+e5f1.get()+"','"+correctanswerval.get()+"')"
	try:
		cursor.execute(sql)
		db.commit()
		#rb1f1.deselect()
		#subjectvalinf1.deselect()
		list1f2.delete(0,END)
		list1f2.insert(END,"Question:"+e1f1.get())
		list1f2.insert(END,"Stored Successfully")
		e1f1.delete(0,END)
		e2f1.delete(0,END)
		e3f1.delete(0,END)
		e4f1.delete(0,END)
		e5f1.delete(0,END)
		rb1f1.deselect()
		rb2f1.deselect()
		rb3f1.deselect()
		rb4f1.deselect()
		rb5f1.deselect()
		rb6f1.deselect()
		rb7f1.deselect()
		rb8f1.deselect()
		rb9f1.deselect()
		#print "Stored Successfully";
	except Exception:
		db.rollback()
		
		list1f2.delete(0,END)
		list1f2.insert(END,"Error Occurred in storing to Database. Please Try Again")
		#print "Error Occurred during saving the question. Please try again"
	db.close()

def generatePaper():
	#
	global generatedQPaper
	db=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	cursor=db.cursor()
	subkeys=subjectQuesCount.keys()
	cnt=0
	for subs in subkeys:
		if(cnt!=0):
			sql=sql+" UNION ALL (SELECT * FROM qBank WHERE subject='"+subs+"' ORDER BY RAND() LIMIT "+subjectQuesCount[subs]+")"
			cnt+=1
		else:
			sql="(SELECT * FROM qBank WHERE subject='"+subs+"' ORDER BY RAND() LIMIT "+subjectQuesCount[subs]+")"
			cnt+=1
	#sql="(SELECT * FROM qBank WHERE subject='"+subject1+"' ORDER BY RAND() LIMIT "+e2f3.get()+") UNION ALL (SELECT * FROM qBank WHERE subject='"+subject2+"' ORDER BY RAND() LIMIT "+e3f3.get()+") UNION ALL (SELECT * FROM qBank WHERE subject='"+subject3+"' ORDER BY RAND() LIMIT "+e4f3.get()+") UNION ALL(SELECT * FROM qBank WHERE subject='"+subject4+"' ORDER BY RAND() LIMIT "+e5f3.get()+") UNION ALL (SELECT * FROM qBank WHERE subject='"+subject5+"' ORDER BY RAND() LIMIT "+e6f3.get()+")";
	print sql
	cursor.execute(sql)
	generatedQPaper=cursor.fetchall()
	db.close()
	list1f2.delete(0,END)
	for q in generatedQPaper:
		list1f2.insert(END,q[1])
	#displayCode
	#db.close()

def evaluateAnswers():
	#tmep=0
	qPaperName=questionpaper1inf3.get()
	db=MySQLdb.connect("127.0.0.1","root","root","pythondb2",3306)
	cursor=db.cursor()
	sql="SELECT qId FROM `"+qPaperName+"`"
	cursor.execute(sql)
	
	qIds=cursor.fetchall()
	db2=MySQLdb.connect("127.0.0.1","root","root","pythonanswers",3306)
	cursor2=db2.cursor()
	sql="SELECT * FROM `"+qPaperName+"`"
	cursor2.execute(sql)
	answerTable=cursor2.fetchall()
	db2.close()
	db.close()
	db=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	cursor=db.cursor()
	answer={}
	for qId in qIds:
		sql="SELECT answer FROM qBank WHERE qId="+str(qId[0])
		cursor.execute(sql)
		# if(cursor.rowcount!=1):
		# 	print "Error!"
		result=cursor.fetchone();
		answer[""+str(qId[0])]=str(result[0])
	#print answer
	score={}
	for row in answerTable:
		user=row[0]
		selection=row[1]
		selection=selection.split(";")
		selCount=len(selection)
		individualScore=0
		print user
		for i in range(0,selCount):
			solutionset=selection[i]
			ans=solutionset.split(",")
			print ans
			print answer[ans[0]]
			if(ans[1]==answer[ans[0]]):
				individualScore=individualScore+1
		score[user]=individualScore
	#tableName="evaluation_"+time.strftime("%d-%m-%Y_%H:%M:%S")
	tableName="evaluation_"+qPaperName
	DeleteTableSQL="DROP TABLE IF EXISTS "+tableName
	cursor.execute(DeleteTableSQL)
	createTableSQL="CREATE TABLE `"+tableName+"` (rollno varchar(9), score int, PRIMARY KEY (rollno))"
	keys=score.keys()
	list1f2.delete(0,END)
	cursor.execute(createTableSQL)	
	for student in keys:
		InsertScoreSQL="INSERT INTO `"+tableName+"` (rollno,score) VALUES('"+student+"',"+str(score[student])+")"
		try:
			cursor.execute(InsertScoreSQL)
			db.commit()
			list1f2.insert(END,student+": "+str(score[student]))

		except Exception:
			db.rollback()
	db.close()



def showQuestions():
	db=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	cursor=db.cursor()
	ShowQuestionSQL="SELECT * FROM qBank"
	val=subjectvalinf4.get();
	if(val!='*'):
		ShowQuestionSQL=ShowQuestionSQL+" WHERE subject='"+val+"'";
	cursor.execute(ShowQuestionSQL)
	questions=cursor.fetchall()
	#display code
	list1f2.delete(0,END)
	for q in questions:
		list1f2.insert(END, q[1])
	db.close()


def setCount():
	count=e2f3.get()
	subName=subjectvalinf3.get()
	subjectQuesCount[""+subName]=count
	subkeys=subjectQuesCount.keys()
	list1f2.delete(0,END)
	for sub in subkeys:
		list1f2.insert(END,sub+":"+ subjectQuesCount[sub])

def addSub():
	db=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	cursor=db.cursor()
	newSub=e1l4.get()
	AddSubjectSQL="INSERT INTO subjects (subject,userId) VALUES('"+newSub+"',1)"
	try:
		cursor.execute(AddSubjectSQL)
		db.commit()
		list1f2.delete(0,END)
		list1f2.insert(END, "Subject Successfully Added")
		om1f1['menu'].delete(0,'end')
		om2f3['menu'].delete(0,'end')
		om1f4['menu'].delete(0,'end')
		om2f4['menu'].delete(0,'end')
		FetchSubject()
		for choice in subjectList:
			om1f1['menu'].add_command(label=choice, command=Tkinter._setit(subjectvalinf1, choice))
			om2f3['menu'].add_command(label=choice, command=Tkinter._setit(subjectvalinf3, choice))
			om1f4['menu'].add_command(label=choice, command=Tkinter._setit(subjectvalinf4, choice))
			om2f4['menu'].add_command(label=choice, command=Tkinter._setit(deletesubjectvalinf4, choice))
		om1f1.update()
		om2f3.update()
		om1f4.update()
		om2f4.update()
		e1l4.delete(0,END)
		db.close()
	except Exception:
		db.rollback()
		db.close()
		


def deleteSub():
	db=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	cursor=db.cursor()
	delSub=deletesubjectvalinf4.get()
	AddSubjectSQL="DELETE FROM subjects WHERE subject='"+delSub+"'"
	try:
		cursor.execute(AddSubjectSQL)
		db.commit()
		list1f2.delete(0,END)
		list1f2.insert(END, "Subject Successfully Deleted")
		om1f1['menu'].delete(0,'end')
		om2f3['menu'].delete(0,'end')
		om1f4['menu'].delete(0,'end')
		om2f4['menu'].delete(0,'end')
		FetchSubject()
		for choice in subjectList:
			om1f1['menu'].add_command(label=choice, command=Tkinter._setit(subjectvalinf1, choice))
			om2f3['menu'].add_command(label=choice, command=Tkinter._setit(subjectvalinf3, choice))
			om1f4['menu'].add_command(label=choice, command=Tkinter._setit(subjectvalinf4, choice))
			om2f4['menu'].add_command(label=choice, command=Tkinter._setit(deletesubjectvalinf4, choice))

		om1f1.update()
		om2f3.update()
		om1f4.update()
		om2f4.update()
		db.close()
	except Exception:
		db.rollback()
		db.close()

qPaperName=""
def commitPaper():
	global qPaperName
	db=MySQLdb.connect("127.0.0.1","root","root","pythondb2",3306)
	db2=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	cursor=db.cursor()
	cursor2=db2.cursor()
	qPaperName=e1f3.get()
	examTime=entrytimer.get()
	if(examTime==""):
		list1f2.delete(0,END)
		list1f2.insert(END, "Please Enter the examination time");
	else:	
		if(qPaperName==""):
			list1f2.delete(0,END)
			list1f2.insert(END, "Please Enter a Name for the Question Paper");
		else:
			try:
				UpdateQPaperListSQL="INSERT INTO qPaperList (papername,timeMin) VALUES('"+qPaperName+"','"+examTime+"')"
				cursor.execute(UpdateQPaperListSQL)
				cursor2.execute(UpdateQPaperListSQL)
				CreateTableSQL="CREATE TABLE `"+qPaperName+"` (qId int,question text,opt1 text,opt2 text,opt3 text,opt4 text, PRIMARY KEY(qId))"
				print CreateTableSQL
				#cursor2.execute(CreateTableSQL)
				#TruncateQuestionPaperSQL="TRUNCATE TABLE quesPaper"
				cursor.execute(CreateTableSQL)
				document=Document()
				document.add_heading(qPaperName, 0)
				p = document.add_paragraph()
				p.alignment = WD_ALIGN_PARAGRAPH.CENTER
				p.add_run('Answer the following Questions:')
				print "Successfully added paragraph"
				p2 = document.add_paragraph()
				docQCount=1
				for row in generatedQPaper:
					sql2="INSERT INTO `"+qPaperName+"` (`qId`,`question`,`opt1`,`opt2`,`opt3`,`opt4`) VALUES('" +str(row[0])+ "','" +row[1]+ "','" +row[2]+ "','" +row[3]+ "','" +row[4]+ "','" +row[5]+ "')"
					#sql3="INSERT INTO `"+qPaperName+"` (`qId`,`question`,`opt1`,`opt2`,`opt3`,`opt4`) VALUES('" +str(row[0])+ "','" +row[1]+ "','" +row[2]+ "','" +row[3]+ "','" +row[4]+ "','" +row[5]+ "')"
					p2.add_run(str(docQCount)+str(row[1])+'\n')
					p2.add_run(' '+'(A)'+str(row[2])+'\n'+' (B)'+str(row[3])+'\n'+' (C)'+str(row[4])+'\n'+' (D)'+str(row[5])+'\n')
					cursor.execute(sql2)
					#cursor2.execute(sql2)
					docQCount+=1
				document.save(qPaperName+'.docx');
				e1f3.delete(0,END)
				e2f3.delete(0,END)
				entrytimer.delete(0,END)
				list1f2.delete(0,END)
				list1f2.insert(END, "Question Paper Committed to Student Database");
				db.commit()
				db2.commit()
				db.close()
				db2.close()
			except Exception,e:
				raise e
				db.rollback()
				db2.rollback()
				list1f2.delete(0,END)
				list1f2.insert(END, "Error in Committing the Question Paper")
				db.close()
				db2.close()
					

	
def FetchSubject():
	global subjectList
	FetchSubListSQL="SELECT * from subjects where userId=1"
	sdb=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	scursor=sdb.cursor()
	scursor.execute(FetchSubListSQL)
	sublist=scursor.fetchall()
	subjectList=[]
	if(scursor.rowcount==0):
		subjectList=["No Subjects Found"]
	else:
		for sub in sublist:
			subjectList.append(sub[2])
	sdb.close()

def FetchQPaper():
	global optList
	FetchQPaperListSQL="SELECT * from qPaperList"
	sdb=MySQLdb.connect("127.0.0.1","root","root","pythontest",3306)
	scursor=sdb.cursor()
	scursor.execute(FetchQPaperListSQL)
	paperlist=scursor.fetchall()
	optList=[]
	if(scursor.rowcount==0):
		optList=["No Question Papers Found"]
	else:
		for pName in paperlist:
			optList.append(pName[0])
	sdb.close()



#-----------------------
scrw=root.winfo_screenwidth()
print scrw

scrh=root.winfo_screenheight()

print scrh
scrndimn=str(scrw)+'x'+str(scrh)

#root.geometry(scrndimn)
root.geometry("1366x768")

subjectvalinf1=StringVar()
subjectvalinf3=StringVar()
subjectvalinf4=StringVar()
deletesubjectvalinf4=StringVar()

questionpaper1inf3=StringVar()
#questionpaper2inf3=StringVar()
FetchSubject()
FetchQPaper()
correctanswerval=StringVar()

labelframe1=LabelFrame(root,highlightbackground="#1E90FF",bg="#D3D3D3",text="Enter Question Details:",height=400,width=1000)
labelframe1.grid(row=0,sticky=W , padx=10, pady=10 , ipady=5 , ipadx=10,columnspan=10)

l1f1=Label(labelframe1 ,bg="#D3D3D3", text="Enter the Question :" , height = 2 , width = 15)
l1f1.grid(row=0 , column = 0 ,pady=10 ,sticky=W, padx=12)

e1f1=Entry(labelframe1 ,highlightbackground="#1E90FF", width = 90)
e1f1.grid(row=0 , column=1 , sticky=W,columnspan=4)

l2f1=Label(labelframe1,bg="#D3D3D3",text="Enter the Options :")
l2f1.grid(row=1 , column=0,sticky=W, padx=12,pady=5)

l3f1=Label(labelframe1,bg="#D3D3D3",text="Option A ------------->")
l3f1.grid(row=2 , column=0 , sticky=W , padx=12,pady=5)

e2f1=Entry(labelframe1,highlightbackground="#1E90FF",width = 50)
e2f1.grid(row=2 , column=1 , sticky=W,pady=5)

l4f1=Label(labelframe1,bg="#D3D3D3",text="Option B ------------->")
l4f1.grid(row=3 , column=0 , sticky=W , padx=12,pady=5)

e3f1=Entry(labelframe1,highlightbackground="#1E90FF",width = 50)
e3f1.grid(row=3 , column=1 , sticky=W,pady=5)

l5f1=Label(labelframe1,bg="#D3D3D3",text="Option C ------------->")
l5f1.grid(row=4 , column=0 , sticky=W , padx=12,pady=5)

e4f1=Entry(labelframe1,highlightbackground="#1E90FF",width = 50)
e4f1.grid(row=4 , column=1 , sticky=W,pady=5)

l6f1=Label(labelframe1,bg="#D3D3D3",text="Option D ------------->")
l6f1.grid(row=5 , column=0 , sticky=W , padx=12,pady=5)

e5f1=Entry(labelframe1,highlightbackground="#1E90FF",width = 50)
e5f1.grid(row=5 , column=1 ,  sticky=W,pady=5)

l7f1=Label(labelframe1,bg="#D3D3D3",text="Choose Subject :")
l7f1.grid(row=7 , column = 0 , padx=12 , sticky=W,pady=5)

om1f1 = OptionMenu(labelframe1, subjectvalinf1, *subjectList)
om1f1.config(width=15 , activebackground="#1E90FF")
om1f1.grid(row = 7 , column =1 ,sticky=W)

l8f1=Label(labelframe1,bg="#D3D3D3",text="Choose the correct answer ->")
l8f1.grid(row=6 , column=0 , padx=12,pady=5)

rb6f1=Radiobutton(labelframe1,activebackground="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",text="A",variable=correctanswerval,value=1)
rb6f1.grid(row=6,column=1,sticky=W,pady=5)

rb7f1=Radiobutton(labelframe1,activebackground="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",text="B",variable=correctanswerval,value=2)
rb7f1.grid(row=6,column=1,sticky=W,padx=40,pady=5)

rb8f1=Radiobutton(labelframe1,activebackground="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",text="C",variable=correctanswerval,value=3)
rb8f1.grid(row=6,column=1,sticky=W,padx=80,pady=5)

rb9f1=Radiobutton(labelframe1,activebackground="#87CEFA",highlightbackground="#D3D3D3",bg="#D3D3D3",text="D",variable=correctanswerval,value=4)
rb9f1.grid(row=6,column=1,sticky=W,padx=120,pady=5)

submitbutton=Button(labelframe1,text="SUBMIT",activebackground="#1E90FF",highlightbackground="#1E90FF",command=insertQuestion , width=35 , bg="#87CEFA")
submitbutton.grid(row=12 ,column=1 ,sticky=W  , pady = 10)

#------------------------------------------------------------------------

#FRAME 2. Displaying Questions

labelframe2=LabelFrame(root,highlightbackground="#1E90FF",bg="#D3D3D3",text="Questions in the database :",height=250,width=860)
labelframe2.grid(row=13 ,sticky=W, padx=10, pady=10 , ipady=5, ipadx=10 , columnspan=10)

scrollbarlb1=Scrollbar(labelframe2)
#list1f2=Listbox(labelframe2,yscrollcommand=scrollbarlb1.set)
#list1f2.grid(row=14,column=0)
list1f2=Listbox(labelframe2,yscrollcommand=scrollbarlb1.set,height=17,width=113)
list1f2.grid(row=14 ,column=0, pady=7 , padx=5)
scrollbarlb1.grid(row=14,column=1,sticky=E+N+S)
scrollbarlb1.config(command=list1f2.yview)
# for line in range(100):
#    list1f2.insert(END, "This is line number " + str(line))

#-----------------------------------------------

#FRAME 3. FUNCTIONS
labelframe3 = LabelFrame(root,highlightbackground="#1E90FF",bg="#D3D3D3",text="Commands :",height=300,width=270)
labelframe3.grid(row=0,column=10 , sticky=N+S+E ,padx=10, pady=10 , rowspan=10 , columnspan =2 )

# templf3 = Label(labelframe3 , text="",highlightbackground="#1E90FF",bg="#D3D3D3")
# templf3.grid(row=1,column=10)

l1f3 = Label(labelframe3 , text="Enter the name of the Question Paper :", highlightbackground="#1E90FF",bg="#D3D3D3")
l1f3.grid(row = 2 , column=10,padx=15, pady=4,sticky=W  )

e1f3 = Entry(labelframe3 , highlightbackground="#1E90FF",width=33 )
e1f3.grid(row=3 , column = 10 ,padx=15, pady=4,sticky=W)

setbutton=Button(labelframe3,text="SET",activebackground="#1E90FF",highlightbackground="#1E90FF",command=setCount , width=12, bg="#87CEFA")
setbutton.grid(row = 9 , sticky=W,column = 10 , padx=15, pady=4)

l3f3 = Label(labelframe3 , text="Exam Duration in minutes :", highlightbackground="#1E90FF",bg="#D3D3D3")
l3f3.grid(row = 4 , column=10 , sticky=W , padx=15, pady=4)

l4f3 = Label(labelframe3 , text="Choose Subject and No. of Questions:", highlightbackground="#1E90FF",bg="#D3D3D3")
l4f3.grid(row = 5 , column=10 , sticky=W , padx=15, pady=4)

entrytimer=Entry(labelframe3 , highlightbackground="#1E90FF", width = 5 )
entrytimer.grid(row = 4 , column=10 , sticky=E , padx=15, pady=4)

om2f3 = OptionMenu(labelframe3, subjectvalinf3, *subjectList)
om2f3.config(width=22 , activebackground="#1E90FF")
om2f3.grid(row = 6 , column =10 ,sticky=W , padx=15, pady=4)

e2f3 = Spinbox(labelframe3, from_=0, to=10 , highlightbackground="#1E90FF" , width=2)
e2f3.grid(row=6 , column = 10 ,sticky=E , padx=15, pady=4)

generatebutton=Button(labelframe3,text="GENERATE",activebackground="#1E90FF",highlightbackground="#1E90FF",command=generatePaper , width=12, bg="#87CEFA")
generatebutton.grid(row = 9 , column =10 , padx=15, pady=4, sticky=E)

acceptButton=Button(labelframe3,text="ACCEPT",activebackground="#1E90FF",highlightbackground="#1E90FF",command=commitPaper , width=30, bg="#87CEFA")
acceptButton.grid(row = 10 , column =10 , padx=15, pady=4 )

l2f3 = Label(labelframe3 , text="Choose Question Paper to Evaluate :", highlightbackground="#1E90FF",bg="#D3D3D3")
l2f3.grid(row = 12 , column=10,padx=15, pady=4 )

#print optList
om1f3 = OptionMenu(labelframe3, questionpaper1inf3, *optList)
om1f3.config(width=15 , activebackground="#1E90FF")
om1f3.grid(row = 13 , column =10 ,padx=15, pady=4)

evaluatebutton=Button(labelframe3,text="EVALUATE",activebackground="#1E90FF",highlightbackground="#1E90FF",command=evaluateAnswers , width=30, bg="#87CEFA")
evaluatebutton.grid(row=14 , column=10 ,padx=15, pady=4)

#----------------------------------------------------

#FRAME 4. DISPLAY QUESTIONS
labelframe4 = LabelFrame(root,highlightbackground="#1E90FF",bg="#D3D3D3",text="Question Bank Access :",height=300,width=320)
labelframe4.grid(row = 13 , column = 10 , pady = 10 , padx = 10 , sticky = N+S+E+W )

l1f4 = Label(labelframe4,text="Choose Subject To View Questions:", highlightbackground="#1E90FF",bg="#D3D3D3")
l1f4.grid(row = 14 , column = 10 , sticky=W , padx=10 , pady=7)

om1f4 = OptionMenu(labelframe4, subjectvalinf4, *subjectList)
om1f4.config(width=14 , activebackground="#1E90FF")
om1f4.grid(row = 15 , column =10 ,sticky=W , padx = 15 , pady=7)

showbutton = Button(labelframe4,text="SHOW",activebackground="#1E90FF",highlightbackground="#1E90FF",command=showQuestions , width=9, bg="#87CEFA")
showbutton.grid(row=15 , column = 10 ,padx=15, pady = 7,sticky=E)

l2f4 = Label(labelframe4,text="Enter/Choose Subject To Add/Delete:", highlightbackground="#1E90FF",bg="#D3D3D3")
l2f4.grid(row = 22 , column = 10 , sticky=W , padx=10 , pady=7)

e1l4=Entry(labelframe4 , highlightbackground="#1E90FF", width = 18 )
e1l4.grid(row = 23, column=10 , sticky=W ,padx=15, pady=7 )

addButton = Button(labelframe4,text="Add Subject",activebackground="#1E90FF",highlightbackground="#1E90FF", command=addSub,width=9, bg="#87CEFA")
addButton.grid(row=23 , column = 10 ,sticky=E ,padx=15, pady=7)

om2f4 = OptionMenu(labelframe4, deletesubjectvalinf4, *subjectList)
om2f4.config(width=14 , activebackground="#1E90FF")
om2f4.grid(row = 24 , column =10 ,sticky=W , padx = 15 , pady=7)

deleteSubbutton = Button(labelframe4,text="Delete Subject",activebackground="#1E90FF",highlightbackground="#1E90FF",command=deleteSub , width=9, bg="#87CEFA")
deleteSubbutton.grid(row=24 , column = 10 , padx=15, pady = 7,sticky=E)

l3f4 = Label(labelframe4,text="Search And Edit Saved Questions:", highlightbackground="#1E90FF",bg="#D3D3D3")
l3f4.grid(row = 25 , column = 10 , sticky=W , padx=10 , pady=7)

editButton = Button(labelframe4,text="EDIT",activebackground="#1E90FF",highlightbackground="#1E90FF", command=editQuestions,width=30, bg="#87CEFA")
editButton.grid(row=26 , column = 10 , padx=15,pady=7 )
#EOF

root.mainloop()